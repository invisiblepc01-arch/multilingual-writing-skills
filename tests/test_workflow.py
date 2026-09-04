import json
import subprocess
import sys
import zipfile
from pathlib import Path

from lxml import etree
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
SCRIPTS = ROOT / "scripts"
W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS = {"w": W}


def run(*args):
    return subprocess.run(
        [sys.executable, *map(str, args)],
        check=True,
        capture_output=True,
        text=True,
    )


def test_fixture_harden_and_audit(tmp_path):
    raw = tmp_path / "raw.docx"
    hardened = tmp_path / "hardened.docx"
    paragraph_report = tmp_path / "paragraph-audit.json"
    run_report = tmp_path / "run-audit.json"

    run(SCRIPTS / "make_bidi_fixture.py", raw)
    run(SCRIPTS / "harden_docx_bidi.py", raw, hardened, "--mode", "auto")
    run(SCRIPTS / "audit_docx_bidi.py", hardened, "--json", paragraph_report)
    run(SCRIPTS / "audit_docx_run_props.py", hardened, "--report", run_report)

    paragraph_result = json.loads(paragraph_report.read_text(encoding="utf-8"))
    assert paragraph_result["passed"] is True
    assert paragraph_result["errors"] == []
    assert paragraph_result["warnings"] == []
    assert paragraph_result["stats"]["paragraphs"] >= 10

    run_result = json.loads(run_report.read_text(encoding="utf-8"))
    assert run_result["passed"] is True
    assert run_result["errors"] == []
    assert run_result["warnings"] == []
    assert run_result["stats"]["visible_runs"] >= 20

    with zipfile.ZipFile(hardened) as archive:
        document = etree.fromstring(archive.read("word/document.xml"))
        styles = etree.fromstring(archive.read("word/styles.xml"))

    assert not styles.xpath(
        './/w:style[@w:styleId="Title"]/w:pPr/w:pBdr', namespaces=NS
    )
    centered = document.xpath(
        './/w:p[contains(string(.), "نباید وسط‌چین")]', namespaces=NS
    )[0]
    assert centered.find("w:pPr/w:jc", NS).get(f"{{{W}}}val") == "right"
    assert centered.find("w:pPr/w:bidi", NS).get(f"{{{W}}}val") == "1"

    first_table = document.xpath(".//w:tbl", namespaces=NS)[0]
    assert first_table.find("w:tblPr/w:bidiVisual", NS).get(f"{{{W}}}val") == "1"
    logical_headers = [
        "".join(cell.xpath(".//w:t/text()", namespaces=NS))
        for cell in first_table.xpath("w:tr[1]/w:tc", namespaces=NS)
    ]
    assert logical_headers == ["شماره", "عنوان", "English"]
    header_alignments = first_table.xpath(
        "w:tr[1]/w:tc/w:p/w:pPr/w:jc/@w:val", namespaces=NS
    )
    assert header_alignments == ["center", "center", "center"]
    header_vertical_alignments = first_table.xpath(
        "w:tr[1]/w:tc/w:tcPr/w:vAlign/@w:val", namespaces=NS
    )
    assert header_vertical_alignments == ["center", "center", "center"]

    all_text = "".join(document.xpath(".//w:t/text()", namespaces=NS))
    assert "SAPنام" not in all_text
    assert "Fioriو" not in all_text
    ezafe_runs = [
        "".join(run.xpath(".//w:t/text()", namespaces=NS))
        for run in document.xpath('.//w:p[contains(string(.), "تجربه")]/w:r', namespaces=NS)
    ]
    assert any("ه\u0654" in value for value in ezafe_runs)
    assert all(not value.startswith("\u0654") for value in ezafe_runs)


def test_explicit_precomposed_ezafe_mode(tmp_path):
    raw = tmp_path / "raw.docx"
    hardened = tmp_path / "precomposed.docx"
    run(SCRIPTS / "make_bidi_fixture.py", raw)
    run(
        SCRIPTS / "harden_docx_bidi.py",
        raw,
        hardened,
        "--mode",
        "auto",
        "--ezafe-mode",
        "precomposed",
    )
    with zipfile.ZipFile(hardened) as archive:
        document = etree.fromstring(archive.read("word/document.xml"))
    all_text = "".join(document.xpath(".//w:t/text()", namespaces=NS))
    assert "\u0654" not in all_text
    assert "ۀ" in all_text


def test_hardener_locks_current_toc_tree(tmp_path):
    raw = tmp_path / "toc.docx"
    hardened = tmp_path / "toc-hardened.docx"
    doc = Document()
    paragraph = doc.add_paragraph("فهرست آزمایشی ")

    def field_char(kind):
        run = OxmlElement("w:r")
        char = OxmlElement("w:fldChar")
        char.set(qn("w:fldCharType"), kind)
        run.append(char)
        paragraph._p.append(run)

    def instruction(value):
        run = OxmlElement("w:r")
        text = OxmlElement("w:instrText")
        text.text = value
        run.append(text)
        paragraph._p.append(run)

    field_char("begin")
    instruction(' TOC \\o "1-3" ')
    field_char("separate")
    field_char("begin")
    instruction(" PAGEREF _Toc123 \\h ")
    field_char("separate")
    paragraph.add_run("1")
    field_char("end")
    field_char("end")
    doc.save(raw)

    run(SCRIPTS / "harden_docx_bidi.py", raw, hardened, "--mode", "auto")
    with zipfile.ZipFile(hardened) as archive:
        document = etree.fromstring(archive.read("word/document.xml"))
    locks = document.xpath(
        './/w:fldChar[@w:fldCharType="begin"]/@w:fldLock', namespaces=NS
    )
    assert locks == ["1", "1"]
