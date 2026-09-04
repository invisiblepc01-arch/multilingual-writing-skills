#!/usr/bin/env python3
import argparse
import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

FA = re.compile(r"[\u0600-\u06ff]")
LATIN = re.compile(r"[A-Za-z0-9]")
FA_LETTER_CLASS = r"\u0621-\u063a\u0641-\u064a\u066e-\u06d3\u06fa-\u06ff"
FA_LETTER = re.compile(rf"[{FA_LETTER_CLASS}]")
FA_TO_LATIN_BOUNDARY = re.compile(rf"([{FA_LETTER_CLASS}])(?=[A-Za-z0-9])")
LATIN_TO_FA_BOUNDARY = re.compile(rf"([A-Za-z0-9])(?=[{FA_LETTER_CLASS}])")
FA_FALLBACK = "B Nazanin"
FA_BOLD_FALLBACK = "B Nazanin Bold"
FA_TITLE_FALLBACK = "B Titr Bold"
EN_FALLBACK = "Arial"
COMBINING_HAMZA_ABOVE = "\u0654"
HEH = "\u0647"
HEH_WITH_YEH_ABOVE = "\u06c0"


def strong_direction(char):
    if FA.match(char):
        return "rtl"
    if char.isascii() and char.isalnum():
        return "ltr"
    return None


def directional_chunks(text, default="rtl"):
    current = next((strong_direction(char) for char in text if strong_direction(char)), default)
    chunks = []
    buffer = ""
    for char in text:
        direction = strong_direction(char)
        if direction is not None and direction != current:
            if buffer:
                chunks.append([current, buffer])
            buffer = char
            current = direction
        else:
            buffer += char
    if buffer:
        chunks.append([current, buffer])
    # In an RTL paragraph, trailing whitespace on an LTR boundary run renders
    # on the wrong side. Move it to the following RTL run.
    for index in range(len(chunks) - 1):
        direction, value = chunks[index]
        next_direction, next_value = chunks[index + 1]
        if direction == "ltr" and next_direction == "rtl":
            match = re.search(r"\s+$", value)
            if match:
                spaces = match.group(0)
                chunks[index][1] = value[: -len(spaces)]
                chunks[index + 1][1] = spaces + next_value
    return [(direction, value) for direction, value in chunks if value]


def split_mixed_text_run(run_el, paragraph_rtl):
    children = list(run_el)
    content_children = [child for child in children if child.tag != qn("w:rPr")]
    if not content_children or any(child.tag != qn("w:t") for child in content_children):
        return
    value = "".join(child.text or "" for child in content_children)
    if not (FA.search(value) and LATIN.search(value)):
        return
    chunks = directional_chunks(value, default="rtl" if paragraph_rtl else "ltr")
    if len(chunks) < 2:
        return
    parent = run_el.getparent()
    if parent is None:
        return
    insert_at = parent.index(run_el)
    parent.remove(run_el)
    for offset, (_, chunk) in enumerate(chunks):
        clone = deepcopy(run_el)
        for child in list(clone):
            if child.tag != qn("w:rPr"):
                clone.remove(child)
        text_node = OxmlElement("w:t")
        if chunk[:1].isspace() or chunk[-1:].isspace():
            text_node.set(qn("xml:space"), "preserve")
        text_node.text = chunk
        clone.append(text_node)
        parent.insert(insert_at + offset, clone)


def add_or_set(parent, tag, value):
    node = parent.find(qn(tag))
    if node is None:
        node = OxmlElement(tag)
        parent.append(node)
    node.set(qn("w:val"), value)
    return node


def text_of(p):
    return "".join(t.text or "" for t in p._p.iter(qn("w:t")))


def normalize_word_ezafe(p, mode="preserve"):
    """Keep heh + combining hamza in one text node/run for stable Word display."""
    text_nodes = list(p._p.iter(qn("w:t")))
    previous = None
    for node in text_nodes:
        value = node.text or ""
        if mode == "precomposed":
            value = value.replace(HEH + COMBINING_HAMZA_ABOVE, HEH_WITH_YEH_ABOVE)
        if value.startswith(COMBINING_HAMZA_ABOVE) and previous is not None:
            previous_value = previous.text or ""
            if previous_value.endswith(HEH):
                previous.text = (
                    previous_value[:-1] + HEH_WITH_YEH_ABOVE
                    if mode == "precomposed"
                    else previous_value + COMBINING_HAMZA_ABOVE
                )
                value = value[1:]
        node.text = value
        if value:
            previous = node


def ensure_script_boundary_spacing(p):
    """Keep a visible separator between adjacent Persian and Latin tokens."""
    previous = None
    for node in p._p.iter(qn("w:t")):
        value = node.text or ""
        value = FA_TO_LATIN_BOUNDARY.sub(r"\1 ", value)
        value = LATIN_TO_FA_BOUNDARY.sub(r"\1 ", value)
        if value and previous is not None:
            previous_value = previous.text or ""
            if previous_value:
                left, right = previous_value[-1], value[0]
                if ((FA_LETTER.match(left) and LATIN.match(right)) or
                        (LATIN.match(left) and FA_LETTER.match(right))):
                    previous.text = previous_value + " "
        node.text = value
        if value:
            previous = node


def set_paragraph(p, rtl, alignment):
    ppr = p._p.get_or_add_pPr()
    add_or_set(ppr, "w:bidi", "1" if rtl else "0")
    add_or_set(ppr, "w:jc", alignment)
    if rtl:
        ind = ppr.find(qn("w:ind"))
        if ind is not None:
            for name in ("left", "start"):
                key = qn(f"w:{name}")
                if key in ind.attrib:
                    del ind.attrib[key]


def set_table_rtl(table):
    tbl_pr = table._tbl.tblPr
    add_or_set(tbl_pr, "w:bidiVisual", "1")
    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), "120")
    indent.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    for old in list(tbl_pr.findall(qn("w:tblBorders"))):
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "D9D9D9")
        borders.append(border)
    tbl_pr.append(borders)


def center_cell_vertically(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    vertical = tc_pr.find(qn("w:vAlign"))
    if vertical is None:
        vertical = OxmlElement("w:vAlign")
        tc_pr.append(vertical)
    vertical.set(qn("w:val"), "center")


def remove_title_borders(doc):
    for style in doc.styles:
        if (style.name or "").casefold() != "title" and (style.style_id or "").casefold() != "title":
            continue
        ppr = style.element.find(qn("w:pPr"))
        if ppr is not None:
            for borders in list(ppr.findall(qn("w:pBdr"))):
                ppr.remove(borders)
    for paragraph in doc.paragraphs:
        if paragraph.style is None or (
            (paragraph.style.name or "").casefold() != "title"
            and (paragraph.style.style_id or "").casefold() != "title"
        ):
            continue
        ppr = paragraph._p.get_or_add_pPr()
        for borders in list(ppr.findall(qn("w:pBdr"))):
            ppr.remove(borders)


def normalize_title_and_heading_colors(doc):
    target_ids = {"title", "subtitle", "heading1", "heading2", "heading3"}
    target_names = {"title", "subtitle", "heading 1", "heading 2", "heading 3"}
    for style in doc.styles:
        if ((style.style_id or "").casefold() not in target_ids and
                (style.name or "").casefold() not in target_names):
            continue
        rpr = style.element.get_or_add_rPr()
        color = rpr.find(qn("w:color"))
        if color is None:
            color = OxmlElement("w:color")
            rpr.append(color)
        color.set(qn("w:val"), "000000")
        for attr in ("themeColor", "themeTint", "themeShade"):
            color.attrib.pop(qn(f"w:{attr}"), None)
    for paragraph in doc.paragraphs:
        style = paragraph.style
        if style is None or ((style.style_id or "").casefold() not in target_ids and
                             (style.name or "").casefold() not in target_names):
            continue
        for run in paragraph._p.iter(qn("w:r")):
            rpr = run.find(qn("w:rPr"))
            if rpr is None:
                rpr = OxmlElement("w:rPr")
                run.insert(0, rpr)
            color = rpr.find(qn("w:color"))
            if color is None:
                color = OxmlElement("w:color")
                rpr.append(color)
            color.set(qn("w:val"), "000000")
            for attr in ("themeColor", "themeTint", "themeShade"):
                color.attrib.pop(qn(f"w:{attr}"), None)


def lock_current_toc_field_tree(doc):
    """Lock the Word-generated TOC and its nested fields after final refresh."""
    stack = []
    for element in doc.element.body.iter():
        if element.tag == qn("w:fldChar"):
            field_type = element.get(qn("w:fldCharType"))
            if field_type == "begin":
                inherited_toc = any(frame["toc"] for frame in stack)
                frame = {"begin": element, "code": "", "toc": inherited_toc}
                if inherited_toc:
                    element.set(qn("w:fldLock"), "1")
                stack.append(frame)
            elif field_type == "separate" and stack:
                frame = stack[-1]
                if re.search(r"(^|\s)TOC(\s|$)", frame["code"], re.IGNORECASE):
                    frame["toc"] = True
                    frame["begin"].set(qn("w:fldLock"), "1")
            elif field_type == "end" and stack:
                stack.pop()
        elif element.tag == qn("w:instrText") and stack:
            stack[-1]["code"] += element.text or ""


def iter_tables(container, seen=None):
    if seen is None:
        seen = set()
    for table in container.tables:
        marker = id(table._tbl)
        if marker in seen:
            continue
        seen.add(marker)
        yield table
        for row in table.rows:
            for cell in row.cells:
                yield from iter_tables(cell, seen)
def style_font_name(paragraph, rtl, run_el=None):
    if not rtl:
        return EN_FALLBACK
    style = paragraph.style
    if style is not None:
        rpr = style.element.find(qn("w:rPr"))
        fonts = None if rpr is None else rpr.find(qn("w:rFonts"))
        if fonts is not None:
            for slot in ("cs", "eastAsia", "hAnsi", "ascii"):
                name = fonts.get(qn(f"w:{slot}"))
                if name:
                    return name
        style_name = (style.name or "").casefold()
        if style_name == "title" or style_name == "heading 1" or "toc title" in style_name:
            return FA_TITLE_FALLBACK
        if style_name.startswith("heading "):
            return FA_BOLD_FALLBACK
    if run_el is not None:
        rpr = run_el.find(qn("w:rPr"))
        if rpr is not None and rpr.find(qn("w:b")) is not None:
            return FA_BOLD_FALLBACK
    return FA_FALLBACK


def set_run(run_el, rtl, font_name):
    rpr = run_el.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run_el.insert(0, rpr)
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.insert(0, fonts)
    for slot in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{slot}"), font_name)
    add_or_set(rpr, "w:rtl", "1" if rtl else "0")
    lang = rpr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rpr.append(lang)
    lang.set(qn("w:val"), "fa-IR" if rtl else "en-US")
    if rtl:
        lang.set(qn("w:bidi"), "fa-IR")
    elif qn("w:bidi") in lang.attrib:
        del lang.attrib[qn("w:bidi")]
    bold = rpr.find(qn("w:b"))
    if rtl and bold is not None:
        bold_cs = rpr.find(qn("w:bCs"))
        if bold_cs is None:
            bold_cs = OxmlElement("w:bCs")
            rpr.append(bold_cs)
        bold_cs.set(qn("w:val"), bold.get(qn("w:val"), "1"))


def current_alignment(p):
    ppr = p._p.get_or_add_pPr()
    jc = ppr.find(qn("w:jc"))
    return None if jc is None else jc.get(qn("w:val"))


def process_paragraph(p, mode, ezafe_mode="preserve", force_alignment=None):
    normalize_word_ezafe(p, ezafe_mode)
    ensure_script_boundary_spacing(p)
    text = text_of(p).strip()
    if not text:
        return
    rtl = mode == "rtl" or (mode == "auto" and bool(FA.search(text)))
    # Persian and mixed paragraphs must be physically right-aligned. Preserve a
    # centered pure-Latin paragraph only; Persian centering requires an explicit
    # force_alignment requested by the caller/user.
    alignment = force_alignment
    if alignment is None and not rtl and mode == "auto" and current_alignment(p) == "center":
        alignment = "center"
    if alignment is None:
        alignment = "right" if rtl else "left"
    set_paragraph(p, rtl, alignment)
    for run in list(p._p.iter(qn("w:r"))):
        split_mixed_text_run(run, rtl)
    for run in p._p.iter(qn("w:r")):
        value = "".join(t.text or "" for t in run.iter(qn("w:t")))
        if not value:
            continue
        run_rtl = rtl and not (LATIN.search(value) and not FA.search(value))
        set_run(run, run_rtl, style_font_name(p, run_rtl, run))


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("input", type=Path)
    ap.add_argument("output", type=Path)
    ap.add_argument("--mode", choices=("auto", "rtl", "ltr"), default="auto")
    ap.add_argument(
        "--ezafe-mode",
        choices=("preserve", "precomposed"),
        default="preserve",
        help="Use precomposed only with explicit approval after Word shows detached U+0654.",
    )
    args = ap.parse_args()
    doc = Document(args.input)
    remove_title_borders(doc)
    normalize_title_and_heading_colors(doc)
    document_has_persian = any(FA.search(text_of(p)) for p in doc.paragraphs)
    if not document_has_persian:
        document_has_persian = any(
            FA.search(text_of(p))
            for table in iter_tables(doc)
            for row in table.rows
            for cell in row.cells
            for p in cell.paragraphs
        )
    for p in doc.paragraphs:
        process_paragraph(p, args.mode, args.ezafe_mode)
    for table in iter_tables(doc):
        if document_has_persian:
            set_table_rtl(table)
        for row_index, row in enumerate(table.rows):
            for cell in row.cells:
                if row_index == 0:
                    center_cell_vertically(cell)
                for p in cell.paragraphs:
                    table_alignment = (
                        "center" if row_index == 0
                        else ("right" if FA.search(text_of(p)) else "left")
                    )
                    process_paragraph(
                        p, args.mode, args.ezafe_mode,
                        force_alignment=table_alignment,
                    )
    seen = set()
    for section in doc.sections:
        for story in (section.header, section.footer):
            if id(story._element) in seen:
                continue
            seen.add(id(story._element))
            for p in story.paragraphs:
                process_paragraph(p, args.mode, args.ezafe_mode)
            for table in iter_tables(story):
                if document_has_persian:
                    set_table_rtl(table)
                for row_index, row in enumerate(table.rows):
                    for cell in row.cells:
                        if row_index == 0:
                            center_cell_vertically(cell)
                        for p in cell.paragraphs:
                            table_alignment = (
                                "center" if row_index == 0
                                else ("right" if FA.search(text_of(p)) else "left")
                            )
                            process_paragraph(
                                p, args.mode, args.ezafe_mode,
                                force_alignment=table_alignment,
                            )
    lock_current_toc_field_tree(doc)
    doc.save(args.output)
    print(args.output)


if __name__ == "__main__":
    main()
