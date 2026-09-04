#!/usr/bin/env python3
import argparse
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("output", type=Path)
    args = ap.parse_args()
    doc = Document()
    title = doc.add_paragraph("عنوان فارسی آزمون", style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style_ppr = doc.styles["Title"].element.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:color"), "4F81BD")
    borders.append(bottom)
    title_style_ppr.append(borders)
    doc.add_heading("۱ - راهنمای آزمایشی Visual Paradigm 15.2", level=1)
    doc.add_paragraph("این یک پاراگراف فارسی شامل BPMN، DFD، عدد 2024 و مسیر C:\\Temp\\file.docx است.")
    centered = doc.add_paragraph("این بند فارسی نباید وسط‌چین باقی بماند")
    centered.alignment = WD_ALIGN_PARAGRAPH.CENTER
    glued = doc.add_paragraph("نمونه SAPنام و FioriوHANA باید مرز خوانا داشته باشد")
    ezafe = doc.add_paragraph()
    ezafe.add_run("تجربه")
    ezafe.add_run("\u0654 پروژه")
    doc.add_paragraph("English paragraph with Persian واژه and punctuation (test).")
    doc.add_heading("۱.۱ - عنوان فارسی با English Span", level=2)
    doc.add_paragraph("مرحله نخست", style="List Number")
    doc.add_paragraph("مرحله دوم با عبارت Save As و کلید Ctrl+S", style="List Number")
    table = doc.add_table(rows=2, cols=3)
    for cell, text in zip(table.rows[0].cells, ["شماره", "عنوان", "English"]):
        cell.text = text
    for cell, text in zip(table.rows[1].cells, ["۱", "نمونه ترکیبی", "BPMN Task"]):
        cell.text = text
    header = doc.sections[0].header.paragraphs[0]
    header.text = "راهنمای فارسی | English Header"
    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "۱"
    doc.save(args.output)
    print(args.output)


if __name__ == "__main__":
    main()
